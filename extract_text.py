import sys
import os

def extract_text_from_pdf(filepath):
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        print(f"---PAGE_COUNT:{len(reader.pages)}---")
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text
    except Exception as e:
        print("---PAGE_COUNT:0---")
        return f"Error reading PDF file: {str(e)}"

def extract_text_from_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        print("---PAGE_COUNT:1---")
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        print("---PAGE_COUNT:0---")
        return f"Error reading DOCX file: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_text.py <filepath>")
        sys.exit(1)
        
    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        sys.exit(1)
        
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        print(extract_text_from_pdf(filepath))
    elif ext == ".docx":
        print(extract_text_from_docx(filepath))
    else:
        # Fallback to plain text
        print("---PAGE_COUNT:1---")
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                print(f.read())
        except Exception as e:
            print(f"Error reading file: {str(e)}")
