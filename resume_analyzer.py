import streamlit as st
import json
import re

# Streamlit Page Config
st.set_page_config(
    page_title="StudentOS — Resume Analyzer & ATS Optimizer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for modern premium dashboard UI
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Space+Grotesk:wght@400;500;600;700&display=swap');
    
    /* Global Styles */
    .stApp {
        font-family: 'Outfit', sans-serif;
    }
    
    h1, h2, h3, .title-text {
        font-family: 'Space Grotesk', sans-serif;
    }
    
    /* Title Banner Styling */
    .header-container {
        padding: 2.5rem;
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.15), rgba(168, 85, 247, 0.15), rgba(236, 72, 153, 0.05));
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 20px;
        margin-bottom: 2rem;
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    
    .glow-effect {
        position: absolute;
        top: -50%;
        left: 50%;
        transform: translateX(-50%);
        width: 600px;
        height: 300px;
        background: radial-gradient(circle, rgba(99, 102, 241, 0.2) 0%, rgba(168, 85, 247, 0) 70%);
        filter: blur(50px);
        z-index: 0;
        pointer-events: none;
    }
    
    .main-title {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(135deg, #6366f1, #a855f7, #ec4899);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.5rem;
        z-index: 1;
        position: relative;
    }
    
    .sub-title {
        font-size: 1.15rem;
        color: #94a3b8;
        max-width: 700px;
        margin: 0 auto;
        z-index: 1;
        position: relative;
        line-height: 1.6;
    }
    
    /* Cards */
    .dashboard-card {
        background-color: rgba(30, 41, 59, 0.5);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 16px;
        padding: 24px;
        margin-bottom: 20px;
    }
    
    /* Badges */
    .badge {
        display: inline-block;
        padding: 6px 12px;
        border-radius: 8px;
        font-size: 0.85rem;
        font-weight: 600;
        margin-right: 8px;
        margin-bottom: 8px;
        font-family: 'Space Grotesk', sans-serif;
    }
    .badge-matching {
        background-color: rgba(34, 197, 94, 0.12);
        color: #4ade80;
        border: 1px solid rgba(34, 197, 94, 0.25);
    }
    .badge-missing {
        background-color: rgba(239, 68, 68, 0.12);
        color: #f87171;
        border: 1px solid rgba(239, 68, 68, 0.25);
    }
    
    /* Rewrite Container */
    .rewrite-card {
        background: rgba(15, 23, 42, 0.4);
        border-left: 4px solid #6366f1;
        border-radius: 0 12px 12px 0;
        padding: 18px;
        margin-bottom: 15px;
        border-top: 1px solid rgba(255, 255, 255, 0.05);
        border-right: 1px solid rgba(255, 255, 255, 0.05);
        border-bottom: 1px solid rgba(255, 255, 255, 0.05);
    }
    
    .original-text {
        color: #ef4444;
        text-decoration: line-through;
        font-size: 0.95rem;
        margin-bottom: 8px;
    }
    
    .optimized-text {
        color: #10b981;
        font-size: 1.05rem;
        font-weight: 600;
        margin-bottom: 8px;
    }
    
    .reason-text {
        color: #94a3b8;
        font-size: 0.88rem;
        font-style: italic;
    }
</style>
""", unsafe_allow_html=True)

# Helper function to extract PDF text
def extract_text_from_pdf(uploaded_file):
    try:
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF file: {str(e)}"

# Helper function to extract DOCX text
def extract_text_from_docx(uploaded_file):
    try:
        from docx import Document
        doc = Document(uploaded_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

# Cleaning Response
def parse_analysis_result(result_text):
    cleaned = result_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r"```$", "", cleaned, flags=re.MULTILINE)
    cleaned = cleaned.strip()
    return json.loads(cleaned)

# Heuristic-based local analyzer fallback
def analyze_resume_heuristics(resume_text, job_description):
    # Normalize texts
    r_words = set(re.findall(r'\b\w+\b', resume_text.lower()))
    jd_words = set(re.findall(r'\b\w+\b', job_description.lower()))
    
    # Common tech/soft skills list to extract
    skills_db = [
        "python", "javascript", "typescript", "react", "vue", "angular", "node", "fastapi", "django", "flask",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "sqlite", "docker", "kubernetes", "aws", "gcp", "azure",
        "git", "github", "ci/cd", "agile", "scrum", "machine learning", "data science", "deep learning", "nlp",
        "html", "css", "tailwind", "bootstrap", "next.js", "vite", "java", "c++", "c#", "go", "rust", "php",
        "communication", "leadership", "problem solving", "teamwork", "analytical", "project management"
    ]
    
    matching_skills = []
    missing_skills = []
    
    for skill in skills_db:
        if re.search(r'\b' + re.escape(skill) + r'\b', job_description.lower()):
            if re.search(r'\b' + re.escape(skill) + r'\b', resume_text.lower()):
                matching_skills.append(skill.title())
            else:
                missing_skills.append(skill.title())
                
    total_skills_in_jd = len(matching_skills) + len(missing_skills)
    if total_skills_in_jd > 0:
        score = int((len(matching_skills) / total_skills_in_jd) * 100)
    else:
        intersection = r_words.intersection(jd_words)
        score = int((len(intersection) / max(len(jd_words), 1)) * 100)
        score = min(max(score, 30), 95)
        
    gap_analysis = []
    if missing_skills:
        gap_analysis.append(f"Add keywords and experience descriptions for missing core skills: {', '.join(missing_skills[:3])}.")
    gap_analysis.append("Incorporate more quantitative results and metrics (e.g., '% improvement', 'reduced latency by X%') to demonstrate impact rather than just listing responsibilities.")
    gap_analysis.append("Optimize the resume structure to place skills in a dedicated, prominent section at the top of your resume.")
    
    # Try to extract bullet points or sentences
    sentences = re.split(r'\.|\n|-', resume_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20 and len(s.strip()) < 120]
    
    rewrites = []
    if sentences:
        target_sentences = sentences[:2]
        for idx, orig in enumerate(target_sentences):
            if idx == 0 and missing_skills:
                skill_to_add = missing_skills[0]
                rw = f"Designed and optimized core modules, integrating {skill_to_add} to enhance application performance and reduce load times by 20%."
                reason = f"Incorporates the highly requested skill '{skill_to_add}' and adds a quantitative result to show direct business impact."
            else:
                rw = f"Spearheaded collaborative development efforts, utilizing industry best practices to deliver project deliverables 15% ahead of schedule."
                reason = "Uses stronger action verbs ('Spearheaded', 'Deliver') and highlights efficiency gains with measurable metrics."
            rewrites.append({
                "original": orig,
                "rewrite": rw,
                "reason": reason
            })
    else:
        rewrites = [
            {
                "original": "Worked on building web pages for the team.",
                "rewrite": "Architected and implemented modular, responsive UI components using React and Tailwind CSS, increasing accessibility scores by 18%.",
                "reason": "Showcases specific front-end tech stack elements from the job description and quantifies layout improvements."
            }
        ]
        
    return {
        "ats_score": score,
        "compatibility_rating": "High Fit" if score >= 80 else ("Medium Fit" if score >= 60 else "Low Fit"),
        "overall_summary": f"Your resume has a matching score of {score}% against the job description. It contains several key matching skills such as {', '.join(matching_skills[:3]) if matching_skills else 'general industry terms'}. However, it is missing some important skills listed in the job description: {', '.join(missing_skills[:3]) if missing_skills else 'none'}. Addressing these gaps and incorporating the suggested bullet rewrites will make your profile significantly stronger for ATS scanners.",
        "matching_skills": matching_skills,
        "missing_skills": missing_skills,
        "gap_analysis": gap_analysis,
        "rewrites": rewrites
    }

# Main Interface Layout

# MAIN CONTAINER
st.markdown("""
<div class="header-container">
    <div class="glow-effect"></div>
    <div class="main-title">Resume Analyzer & ATS Optimizer</div>
    <div class="sub-title">Bridge the gap between your resume and your dream role. Upload your resume, paste the target job description, and get instant score optimization.</div>
</div>
""", unsafe_allow_html=True)

# Layout: Upload & Job Description
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("📄 Upload Resume")
    uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
    
    resume_text = ""
    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()
        if file_extension == "pdf":
            resume_text = extract_text_from_pdf(uploaded_file)
        elif file_extension == "docx":
            resume_text = extract_text_from_docx(uploaded_file)
            
        if resume_text.startswith("Error"):
            st.error(resume_text)
            resume_text = ""
        elif resume_text:
            st.success(f"Successfully loaded: {uploaded_file.name} ({len(resume_text)} characters)")
            
            # Show a brief preview of the resume
            with st.expander("🔍 Preview Resume Text"):
                st.text_area("Extracted Content", resume_text, height=200, disabled=True)

with col2:
    st.subheader("💼 Paste Job Description")
    job_description = st.text_area("Paste the job description of your target role", height=230, placeholder="We are looking for a Software Engineer with experience in Python, FastAPI, Docker...")

# Run Button
st.markdown("<div style='text-align: center; margin: 30px 0;'>", unsafe_allow_html=True)
run_btn = st.button("🚀 Analyze Compatibility & Optimize", use_container_width=True, type="primary")
st.markdown("</div>", unsafe_allow_html=True)

if run_btn:
    if not resume_text:
        st.error("📄 Please upload a valid Resume file (PDF or DOCX).")
    elif not job_description.strip():
        st.error("💼 Please paste a target Job Description to compare against.")
    else:
        with st.spinner("⏳ Analyzing resume alignment against the job description... This may take a few seconds."):
            try:
                data = analyze_resume_heuristics(resume_text, job_description)
                
                # Setup visual score gauge
                score = data.get("ats_score", 0)
                if score >= 80:
                    score_color = "#22c55e" # Green
                    rating_label = "High Match Fit"
                elif score >= 60:
                    score_color = "#eab308" # Yellow
                    rating_label = "Moderate Match Fit"
                else:
                    score_color = "#ef4444" # Red
                    rating_label = "Low Match Fit"
                
                # Store analysis in session state to preserve across interactions
                st.session_state["analysis_result"] = data
                st.session_state["analysis_score_color"] = score_color
                st.session_state["analysis_rating_label"] = rating_label
                
            except Exception as e:
                st.error(f"❌ An error occurred during analysis: {str(e)}")

# Render Results from Session State
if "analysis_result" in st.session_state:
    data = st.session_state["analysis_result"]
    score_color = st.session_state["analysis_score_color"]
    rating_label = st.session_state["analysis_rating_label"]
    score = data.get("ats_score", 0)
    
    st.markdown("<hr style='border-top: 1px solid rgba(255,255,255,0.08); margin: 30px 0;' />", unsafe_allow_html=True)
    st.subheader("📊 Analysis Results")
    
    # Grid: Score Indicator & Match Summary
    res_col1, res_col2 = st.columns([1, 2], gap="medium")
    
    with res_col1:
        st.markdown(f"""
        <div style="background: rgba(30, 41, 59, 0.4); border: 1px solid rgba(255, 255, 255, 0.08); border-radius: 16px; padding: 30px; text-align: center; height: 100%; display: flex; flex-direction: column; justify-content: center; align-items: center;">
            <h3 style="margin: 0; font-size: 1.1rem; color: #94a3b8; font-weight: 500; font-family: 'Space Grotesk';">ATS Match Rating</h3>
            <div style="font-size: 4.5rem; font-weight: 800; color: {score_color}; margin: 15px 0; font-family: 'Space Grotesk'; line-height: 1;">{score}%</div>
            <div style="display: inline-block; padding: 6px 16px; border-radius: 20px; background: {score_color}18; color: {score_color}; border: 1px solid {score_color}35; font-weight: 600; font-size: 0.95rem; font-family: 'Space Grotesk';">
                {rating_label}
            </div>
        </div>
        """, unsafe_allow_html=True)
        
    with res_col2:
        st.markdown("""<div style="height: 100%; display: flex; flex-direction: column; justify-content: center;">""", unsafe_allow_html=True)
        st.write("### 📝 Match Summary")
        st.write(data.get("overall_summary", "No summary provided."))
        st.markdown("</div>", unsafe_allow_html=True)

    # Tabs for detailed breakdown
    tab1, tab2, tab3 = st.tabs(["🎯 Skill & Keyword Match", "🛠️ Gap Analysis", "✨ Tailored Experience Rewrites"])
    
    with tab1:
        st.markdown("### Keyword Alignment Analysis")
        st.write("See how the target keywords, technical skills, and tools in the job description align with your resume.")
        
        tab1_col1, tab1_col2 = st.columns([1, 1], gap="medium")
        
        with tab1_col1:
            st.markdown("""
            <div style="border: 1px solid rgba(34, 197, 94, 0.2); background: rgba(34, 197, 94, 0.04); border-radius: 12px; padding: 20px; height: 100%;">
                <h4 style="color: #4ade80; margin-top: 0; margin-bottom: 15px; font-family: 'Space Grotesk';">✅ Matching Keywords & Skills</h4>
            """, unsafe_allow_html=True)
            
            matching = data.get("matching_skills", [])
            if matching:
                for skill in matching:
                    st.markdown(f'<span class="badge badge-matching">{skill}</span>', unsafe_allow_html=True)
            else:
                st.write("No matching skills identified. Consider adding keywords from the job description.")
            st.markdown("</div>", unsafe_allow_html=True)
            
        with tab1_col2:
            st.markdown("""
            <div style="border: 1px solid rgba(239, 68, 68, 0.2); background: rgba(239, 68, 68, 0.04); border-radius: 12px; padding: 20px; height: 100%;">
                <h4 style="color: #f87171; margin-top: 0; margin-bottom: 15px; font-family: 'Space Grotesk';">❌ Missing Keywords & Skills</h4>
            """, unsafe_allow_html=True)
            
            missing = data.get("missing_skills", [])
            if missing:
                for skill in missing:
                    st.markdown(f'<span class="badge badge-missing">{skill}</span>', unsafe_allow_html=True)
            else:
                st.write("🎉 Excellent! No major missing keywords identified.")
            st.markdown("</div>", unsafe_allow_html=True)

    with tab2:
        st.markdown("### Actionable Gaps & Fixes")
        st.write("Here are the specific formatting issues, structural gaps, or missing contexts to address to score higher.")
        
        gaps = data.get("gap_analysis", [])
        if gaps:
            for i, gap in enumerate(gaps):
                st.markdown(f"""
                <div style="display: flex; gap: 15px; align-items: flex-start; margin-bottom: 12px; background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); border-radius: 10px; padding: 12px 18px;">
                    <div style="font-family: 'Space Grotesk'; font-weight: 700; color: #a855f7; font-size: 1.1rem;">0{i+1}</div>
                    <div style="font-size: 0.95rem; color: #cbd5e1; line-height: 1.5;">{gap}</div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.success("No gaps detected! Your resume structure matches the industry standard.")

    with tab3:
        st.markdown("### Tailored Bullet-by-Bullet Rewrites")
        st.write("Replace weak bullet points with optimized ones that highlight the target keywords and demonstrate quantified impact using active verbs.")
        
        rewrites = data.get("rewrites", [])
        if rewrites:
            for item in rewrites:
                orig = item.get("original", "")
                rw = item.get("rewrite", "")
                reason = item.get("reason", "")
                
                st.markdown(f"""
                <div class="rewrite-card">
                    <div style="font-size: 0.8rem; font-weight: bold; text-transform: uppercase; color: #ef4444; letter-spacing: 0.05em; margin-bottom: 3px;">Original bullet point</div>
                    <div class="original-text">{orig}</div>
                    <div style="font-size: 0.8rem; font-weight: bold; text-transform: uppercase; color: #10b981; letter-spacing: 0.05em; margin-top: 12px; margin-bottom: 3px;">Optimized ATS rewrite</div>
                    <div class="optimized-text">{rw}</div>
                    <div style="margin-top: 10px; border-top: 1px dashed rgba(255, 255, 255, 0.08); padding-top: 8px;">
                        <span style="font-size: 0.8rem; font-weight: bold; text-transform: uppercase; color: #a855f7; letter-spacing: 0.05em; display: block; margin-bottom: 3px;">Strategy & Context</span>
                        <div class="reason-text">{reason}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No bullet rewrites suggested. Your experience statements are already well-written.")
